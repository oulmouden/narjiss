from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import json


ROOT = Path(r"C:\xampp\htdocs\mimosas")
OUT_DIR = Path(r"C:\xampp\htdocs\narjiss\mimosas-report")
ASSET_DIR = OUT_DIR / "assets"
DOCX_OUT = OUT_DIR / "rapport-investisseur-hotel-mimosas.docx"

OUT_DIR.mkdir(exist_ok=True)
ASSET_DIR.mkdir(exist_ok=True)


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


project = load_json(ROOT / "data" / "projects.json")[0]
contacts = load_json(ROOT / "data" / "contacts.json")


PALETTE = {
    "brown": "7A3F13",
    "gold": "B58A37",
    "orange": "C15A00",
    "ink": "1F2937",
    "muted": "6B7280",
    "cream": "F7F1E6",
    "blue": "0F5E7A",
    "pale": "F8FAFC",
    "line": "E5E7EB",
}


def prepare_image(rel_or_abs, name, max_width=1800, max_height=1200, quality=86):
    src = Path(rel_or_abs)
    if not src.is_absolute():
        src = ROOT / src
    if not src.exists():
        return None
    suffix = src.suffix.lower()
    out = ASSET_DIR / f"{name}{'.png' if suffix == '.png' and 'logo' in name else '.jpg'}"
    with Image.open(src) as im:
        im = im.convert("RGBA" if out.suffix == ".png" else "RGB")
        im.thumbnail((max_width, max_height), Image.LANCZOS)
        if out.suffix == ".png":
            im.save(out, optimize=True)
        else:
            im.save(out, quality=quality, optimize=True)
    return out


images = {
    "logo": prepare_image("images/logo-mimosas-tranparent.png", "logo", 900, 700),
    "facade": prepare_image("images/projects/mimosas/devanture-hotel-Mimosas.png", "facade"),
    "entree": prepare_image("images/slider/entree hotel.png", "entree"),
    "reception": prepare_image("images/slider/reception hotel.png", "reception"),
    "chambre": prepare_image("images/slider/chambre 01-01.png", "chambre"),
    "couloir": prepare_image("images/slider/couloir-01.png", "couloir"),
    "rooftop": prepare_image("images/slider/rooftop-01.png", "rooftop"),
    "rooftop_bar": prepare_image("images/slider/rooftop-02.png", "rooftop_bar"),
    "restaurant": prepare_image("images/slider/salle restaurant.png", "restaurant"),
    "sdb": prepare_image("images/slider/sdb-02.png", "sdb"),
    "plan": prepare_image("images/projects/mimosas/floorplan-2.png", "plan", 1400, 1000),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="E5E7EB", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.2)
                    run.font.name = "Aptos"
        if header and r_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, PALETTE["brown"])
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Aptos"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Aptos Display"
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(PALETTE["brown"] if level == 1 else PALETTE["ink"])
    run.font.size = Pt(20 if level == 1 else 13.5)
    return p


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    if bold_start and text.startswith(bold_start):
        add_run(p, bold_start, bold=True, color=PALETTE["ink"], size=10.2)
        add_run(p, text[len(bold_start):], size=10.2, color=PALETTE["ink"])
    else:
        add_run(p, text, size=10.2, color=PALETTE["ink"])
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.55)
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(PALETTE["ink"])
    if not p.runs:
        add_run(p, text, size=10, color=PALETTE["ink"])
    else:
        p.runs[0].text = text
    return p


def add_callout(doc, title, body, fill="F7F1E6"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "D7B56D", "10")
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    add_run(p, title, bold=True, color=PALETTE["brown"], size=11.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    add_run(p2, body, color=PALETTE["ink"], size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_photo(doc, image_path, caption, width=6.2):
    if not image_path:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color=PALETTE["muted"], size=8.6)


def add_photo_grid(doc, items):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx in range(0, len(items), 2):
        row = table.add_row()
        for col in range(2):
            cell = row.cells[col]
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, 80, 80, 140, 80)
            item_idx = idx + col
            if item_idx >= len(items):
                continue
            path, caption = items[item_idx]
            if path:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(path), width=Inches(3.05))
                c = cell.add_paragraph()
                c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(c, caption, italic=True, color=PALETTE["muted"], size=8.2)


def add_kpi_table(doc):
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    vals = [
        ("Adresse", project["location"]["fr"]),
        ("Positionnement", "Hôtel urbain + rooftop + restauration"),
        ("Digital", "Site multilingue, cartes Leaflet, galerie et visite 360"),
        ("POI", f"{project.get('poi_count', 0)} points d'intérêt cartographiés"),
    ]
    for i, (label, value) in enumerate(vals):
        cell = hdr[i]
        set_cell_shading(cell, PALETTE["cream"])
        set_cell_border(cell, "E6D3AA")
        set_cell_margins(cell, 180, 160, 180, 160)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label.upper(), bold=True, color=PALETTE["brown"], size=8.4)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, value, color=PALETTE["ink"], size=9.2)


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if images["logo"]:
    p.add_run().add_picture(str(images["logo"]), width=Inches(2.15))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Hôtel Mimosas")
r.font.name = "Aptos Display"
r.font.size = Pt(30)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(PALETTE["brown"])
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Rapport investisseur, business plan synthétique et argumentaire de cession", color=PALETTE["ink"], size=14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "N°16, avenue Lalla Abla, 85000 Tiznit | Actif hôtelier à finaliser et exploiter", color=PALETTE["muted"], size=10.5)
add_photo(doc, images["facade"], "Devanture et identité visuelle du projet Hôtel Mimosas", width=6.65)
add_kpi_table(doc)
doc.add_page_break()

add_heading(doc, "1. Synthèse exécutive", 1)
add_callout(
    doc,
    "Thèse d'investissement",
    "Mimosas est un actif hôtelier à Tiznit dont la valeur repose sur un potentiel déjà visible : identité de marque, espaces aménagés, rooftop, restaurant, visite 360, galerie photo, site multilingue et localisation urbaine proche des flux de la médina. L'opportunité pour un repreneur est d'acquérir un projet avancé, d'exécuter les dernières finitions, puis de lancer une exploitation orientée séjour, restauration, rooftop et expérience locale.",
)
add_body(doc, "Le projet n'est pas présenté comme une promesse immobilière abstraite, mais comme une plateforme d'exploitation hôtelière. L'acheteur potentiel doit pouvoir se projeter rapidement : finaliser, classer, commercialiser, ouvrir, mesurer, optimiser.")
add_body(doc, "Ce rapport propose une lecture commerciale du dossier : pourquoi l'actif a du potentiel, quels leviers de chiffre d'affaires activer, quelles finitions prioriser et comment présenter la vente à un investisseur hôtelier.")

add_heading(doc, "2. Fiche actif", 1)
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Critère"
table.rows[0].cells[1].text = "Lecture investisseur"
rows = [
    ("Nom commercial", "Hôtel Mimosas"),
    ("Adresse", project["location"]["fr"]),
    ("Coordonnées GPS", f"{project['lat']}, {project['lng']}"),
    ("Concept", "Hôtel urbain avec rooftop, espaces restauration/bar et potentiel séjour court ou étape touristique."),
    ("État du projet", "Actif à potentiel avancé : image de marque, espaces visibles, contenus photos, visite virtuelle et site web déjà préparés. Les travaux, conformité et classement restent à confirmer par audit technique."),
    ("Actifs digitaux", "Site vitrine multilingue, cartographie Leaflet, points d'intérêt, galerie photo, bouton Visite 360 et back-office."),
    ("URL visite 360", project.get("tour_url", "Mimosas/Tour/index.htm")),
]
for left, right in rows:
    cells = table.add_row().cells
    cells[0].text = left
    cells[1].text = right
shade_table(table)

add_heading(doc, "3. Contexte marché et territoire", 1)
add_body(doc, "Le Maroc bénéficie d'une dynamique touristique forte. Les chiffres publics du Ministère du Tourisme indiquent 19,8 millions de visiteurs en 2025, une progression de 14% par rapport à 2024, ainsi que 138 milliards de dirhams de recettes touristiques en 2025.")
add_body(doc, "Tiznit offre une proposition distincte des destinations très saturées : médina, remparts, artisanat d'argent, proximité d'Aglou, accès aux paysages du Souss Massa et positionnement de séjour paisible. Cette différenciation peut servir un hôtel indépendant s'il sait raconter une expérience locale claire.")
add_bullet(doc, "Tiznit est présentée par Visit Agadir comme une destination de la région Souss Massa, avec médina, patrimoine, artisanat et accès à des sites naturels.")
add_bullet(doc, "Aglou, à environ 15 km de Tiznit selon Visit Agadir, renforce le potentiel d'excursions mer, détente, activités et courts séjours.")
add_bullet(doc, "Le rooftop permet de créer une offre visible au-delà de l'hébergement : café, restauration légère, soirées, privatisations et rendez-vous locaux.")

add_heading(doc, "4. Atouts spécifiques de Mimosas", 1)
table = doc.add_table(rows=1, cols=3)
table.rows[0].cells[0].text = "Atout"
table.rows[0].cells[1].text = "Impact commercial"
table.rows[0].cells[2].text = "Action recommandée"
rows = [
    ("Rooftop", "Différenciation forte dans une ville de passage et de découverte.", "Créer une offre sunset, petit-déjeuner, café, tapas, événements privés."),
    ("Restaurant / comptoir", "Revenus additionnels hors nuitées.", "Définir une carte courte, locale, rentable et facile à exécuter."),
    ("Chambres et espaces déjà photogéniques", "Commercialisation rapide possible après finitions.", "Finaliser literie, éclairage, textiles, signalétique et shooting pro."),
    ("Visite 360", "Réduit l'incertitude investisseur et client.", "L'utiliser dans la vente, les annonces, Google Business et le site."),
    ("Site multilingue", "Base de distribution et de crédibilité déjà prête.", "Connecter réservation, WhatsApp, SEO local et réseaux sociaux."),
    ("Carte + POI", "Valorise l'adresse et le territoire.", "Créer des itinéraires clients : médina, Aglou, Tafraout, artisanat."),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
shade_table(table)

doc.add_page_break()
add_heading(doc, "5. Dossier visuel", 1)
add_body(doc, "Les visuels ci-dessous montrent un actif déjà lisible : identité, façade, zones d'accueil, espaces chambres, rooftop et restaurant. Ils doivent être utilisés comme preuves de potentiel dans le dossier de vente.")
add_photo_grid(doc, [
    (images["facade"], "Devanture Hôtel Mimosas"),
    (images["entree"], "Entrée et signalétique"),
    (images["reception"], "Réception / accueil"),
    (images["chambre"], "Chambre témoin"),
    (images["couloir"], "Circulation intérieure"),
    (images["sdb"], "Salle d'eau"),
    (images["rooftop"], "Rooftop et terrasse"),
    (images["restaurant"], "Salle restaurant"),
])

doc.add_page_break()
add_heading(doc, "6. Business plan synthétique", 1)
add_callout(
    doc,
    "Important",
    "Les chiffres précis doivent être finalisés après audit : prix d'acquisition, nombre exact de chambres exploitables, capacité rooftop, CAPEX restant, statut juridique, autorisations, classement, masse salariale et politique tarifaire. Le tableau ci-dessous donne une architecture de business plan à compléter, pas une valorisation définitive.",
    fill="FFF7ED",
)
table = doc.add_table(rows=1, cols=4)
table.rows[0].cells[0].text = "Pilier"
table.rows[0].cells[1].text = "Hypothèse opérationnelle"
table.rows[0].cells[2].text = "KPI à suivre"
table.rows[0].cells[3].text = "Priorité"
rows = [
    ("Hébergement", "Vente directe + OTA + clientèle locale, MRE, couples, voyageurs d'étape, professionnels.", "Taux d'occupation, ADR, RevPAR, avis clients.", "Très haute"),
    ("Rooftop", "Offre café/soft, brunch, sunset, tapas, événements privés.", "Ticket moyen, couverts/jour, marge brute.", "Très haute"),
    ("Restaurant", "Carte courte à rotation rapide, produits locaux, room service simple.", "Coût matière, marge, satisfaction.", "Haute"),
    ("Expériences", "Itinéraires Tiznit, médina, Aglou, artisanat, excursions partenaires.", "Commissions, ventes additionnelles.", "Moyenne"),
    ("Digital", "Site multilingue, Google Business, WhatsApp, visite 360, SEO local.", "Trafic, demandes, conversion directe.", "Très haute"),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
shade_table(table)

add_heading(doc, "Modèle de revenus", 2)
add_bullet(doc, "Nuitées : chambres vendues en direct, OTA, corporate local, week-ends et séjours courts.")
add_bullet(doc, "Restauration : petit-déjeuner, café, rooftop, carte courte, privatisations.")
add_bullet(doc, "Services : navette, excursions, guide local, partenariats avec artisans et activités autour de Tiznit/Aglou.")
add_bullet(doc, "Image de marque : création d'un lieu identifiable, photographiable, partageable et facilement recommandable.")

add_heading(doc, "7. Plan de reprise en 90 jours", 1)
table = doc.add_table(rows=1, cols=4)
table.rows[0].cells[0].text = "Phase"
table.rows[0].cells[1].text = "Objectif"
table.rows[0].cells[2].text = "Actions"
table.rows[0].cells[3].text = "Livrables"
rows = [
    ("Jours 1-15", "Audit et sécurisation", "Audit technique, juridique, conformité, sécurité incendie, licences, inventaire travaux.", "Liste CAPEX, risques, calendrier."),
    ("Jours 16-45", "Finitions et standard hôtelier", "Chambres, literie, éclairage, plomberie, signalétique, rooftop, cuisine, réception.", "Actif prêt pré-ouverture."),
    ("Jours 46-70", "Commercialisation", "Photos finales, Google Business, OTA, site, WhatsApp, visite 360, contenus réseaux.", "Tunnel de réservation."),
    ("Jours 71-90", "Soft opening", "Ouverture progressive, collecte avis, ajustement pricing, formation équipe.", "Premiers revenus et preuve d'exploitation."),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
shade_table(table)

add_heading(doc, "8. Argumentaire de vente", 1)
add_body(doc, "Message court à utiliser auprès des investisseurs :", bold_start="Message court à utiliser auprès des investisseurs :")
add_callout(
    doc,
    "Pitch",
    "Hôtel Mimosas est une opportunité de reprise d'un actif hôtelier à fort potentiel à Tiznit : un projet déjà identifiable, doté d'un rooftop, d'espaces restaurant, d'une identité visuelle, d'un site multilingue, d'une visite virtuelle 360 et d'un storytelling territorial. Le repreneur ne part pas d'une page blanche : il reprend une base avancée, finalise les derniers travaux, structure l'exploitation et capte la valeur d'un hôtel indépendant dans une ville patrimoniale entre médina, Anti-Atlas et littoral d'Aglou.",
)
add_heading(doc, "Arguments clés", 2)
add_bullet(doc, "Actif tangible : photos, espaces et identité déjà visibles.")
add_bullet(doc, "Potentiel d'exploitation multiple : chambres, rooftop, restaurant, événements, expériences locales.")
add_bullet(doc, "Destination à raconter : Tiznit, médina, remparts, artisanat d'argent, Aglou, Souss Massa.")
add_bullet(doc, "Outils commerciaux prêts : site multilingue, cartes, points d'intérêt, galerie et visite 360.")
add_bullet(doc, "Stratégie de reprise claire : finaliser, ouvrir, mesurer, optimiser.")

add_heading(doc, "9. Points de vigilance pour l'acheteur", 1)
add_bullet(doc, "Confirmer le nombre exact de chambres exploitables, surfaces, capacité rooftop et capacité restaurant.")
add_bullet(doc, "Chiffrer le CAPEX restant par lot : sécurité, plomberie, électricité, mobilier, cuisine, enseigne, accessibilité, classement.")
add_bullet(doc, "Vérifier les autorisations, licences, conformité ERP/hôtel, assurance, fiscalité et statut foncier.")
add_bullet(doc, "Établir un budget de lancement : équipe, consommables, outils de réservation, marketing, shooting final.")
add_bullet(doc, "Formaliser un compte d'exploitation prévisionnel en trois scénarios : prudent, cible, ambitieux.")

add_heading(doc, "10. Documents à préparer pour une vente plus forte", 1)
table = doc.add_table(rows=1, cols=3)
table.rows[0].cells[0].text = "Document"
table.rows[0].cells[1].text = "Utilité"
table.rows[0].cells[2].text = "Statut conseillé"
rows = [
    ("Dossier technique", "Rassurer sur travaux, conformité, surfaces.", "À compléter par audit."),
    ("Compte d'exploitation prévisionnel", "Montrer la logique de rentabilité.", "À construire après CAPEX et inventaire chambres."),
    ("Album photos final", "Faire ressentir l'expérience.", "À produire après finitions."),
    ("Visite 360", "Preuve immersive du potentiel.", "Déjà disponible."),
    ("Site web / back-office", "Prouver la base commerciale.", "Déjà préparé."),
    ("Pack investisseur", "Faciliter la décision.", "Ce rapport peut servir de base."),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
shade_table(table)

add_heading(doc, "11. Sources et références", 1)
add_body(doc, "Sources externes consultées pour le contexte marché et territorial :")
sources = [
    "Ministère du Tourisme, chiffres clés 2025 : https://mtaess.gov.ma/fr/chiffres-cles/",
    "Ministère du Tourisme, communiqué du 9 janvier 2026 : https://mtaess.gov.ma/fr/tourisme-le-maroc-atteint-un-seuil-historique-de-pres-de-20-millions-de-touristes-en-2025/",
    "Visit Agadir / Souss Massa, médina de Tiznit : https://visitagadir.com/destinations/medina-de-tiznit/",
    "Visit Agadir / Souss Massa, Aglou : https://visitagadir.com/destinations/aglou-la-plage-desertique/",
    "Données internes Mimosas : site local, galerie, visite 360, data/projects.json, data/contacts.json.",
]
for source in sources:
    add_bullet(doc, source)

add_heading(doc, "Conclusion", 1)
add_body(doc, "La valeur de Mimosas se joue dans la conversion d'un projet avancé en actif exploité. Le repreneur idéal n'achète pas seulement des murs ou des finitions : il achète une adresse, une marque, des espaces à potentiel, un rooftop différenciant, une base digitale et un récit touristique local prêt à être renforcé.")
add_body(doc, "La prochaine étape consiste à transformer ce rapport en dossier de cession chiffré : prix, CAPEX, calendrier, capacité réelle, prévisionnel d'exploitation et modalités de transaction.")

doc.core_properties.title = "Rapport investisseur - Hôtel Mimosas"
doc.core_properties.subject = "Business plan synthétique et argumentaire de vente"
doc.core_properties.author = "Hôtel Mimosas"
doc.save(DOCX_OUT)
print(DOCX_OUT)
