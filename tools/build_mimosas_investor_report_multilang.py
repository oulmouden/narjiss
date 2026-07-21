from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import json


ROOT = Path(r"C:\xampp\htdocs\mimosas")
OUT_DIR = Path(r"C:\xampp\htdocs\narjiss\mimosas-report\multilang")
ASSET_DIR = OUT_DIR / "assets"
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)

project = json.loads((ROOT / "data" / "projects.json").read_text(encoding="utf-8"))[0]

PALETTE = {
    "brown": "7A3F13",
    "gold": "B58A37",
    "orange": "C15A00",
    "ink": "1F2937",
    "muted": "6B7280",
    "cream": "F7F1E6",
    "line": "E5E7EB",
}


def prepare_image(rel_or_abs, name, max_width=1800, max_height=1200, quality=86):
    src = Path(rel_or_abs)
    if not src.is_absolute():
        src = ROOT / src
    if not src.exists():
        return None
    out = ASSET_DIR / f"{name}{'.png' if 'logo' in name else '.jpg'}"
    with Image.open(src) as im:
        im = im.convert("RGBA" if out.suffix == ".png" else "RGB")
        im.thumbnail((max_width, max_height), Image.LANCZOS)
        if out.suffix == ".png":
            im.save(out, optimize=True)
        else:
            im.save(out, quality=quality, optimize=True)
    return out


IMAGES = {
    "logo": prepare_image("images/logo-mimosas-tranparent.png", "logo", 900, 700),
    "facade": prepare_image("images/projects/mimosas/devanture-hotel-Mimosas.png", "facade"),
    "entree": prepare_image("images/slider/entree hotel.png", "entree"),
    "reception": prepare_image("images/slider/reception hotel.png", "reception"),
    "chambre": prepare_image("images/slider/chambre 01-01.png", "chambre"),
    "couloir": prepare_image("images/slider/couloir-01.png", "couloir"),
    "rooftop": prepare_image("images/slider/rooftop-01.png", "rooftop"),
    "restaurant": prepare_image("images/slider/salle restaurant.png", "restaurant"),
    "sdb": prepare_image("images/slider/sdb-02.png", "sdb"),
}


def tr_text():
    address = {
        "fr": project["location"]["fr"],
        "en": project["location"]["en"],
        "ar": project["location"]["ar"],
        "es": project["location"]["es"],
    }
    return {
        "fr": {
            "title": "Hôtel Mimosas",
            "subtitle": "Rapport investisseur, business plan synthétique et argumentaire de cession",
            "meta": f"{address['fr']} | Actif hôtelier à finaliser et exploiter",
            "caption_cover": "Devanture et identité visuelle du projet Hôtel Mimosas",
            "labels": ["Adresse", "Positionnement", "Digital", "POI"],
            "kpis": [address["fr"], "Hôtel urbain + rooftop + restauration", "Site multilingue, cartes Leaflet, galerie et visite 360", f"{project.get('poi_count', 0)} points d'intérêt cartographiés"],
            "sections": [
                ("1. Synthèse exécutive", [
                    ("callout", "Thèse d'investissement", "Mimosas est un actif hôtelier à Tiznit dont la valeur repose sur un potentiel déjà visible : identité de marque, espaces aménagés, rooftop, restaurant, visite 360, galerie photo, site multilingue et localisation urbaine proche des flux de la médina. L'opportunité pour un repreneur est d'acquérir un projet avancé, d'exécuter les dernières finitions, puis de lancer une exploitation orientée séjour, restauration, rooftop et expérience locale."),
                    ("p", "Le projet n'est pas présenté comme une promesse immobilière abstraite, mais comme une plateforme d'exploitation hôtelière. L'acheteur potentiel doit pouvoir se projeter rapidement : finaliser, classer, commercialiser, ouvrir, mesurer, optimiser."),
                ]),
                ("2. Fiche actif", "asset_table"),
                ("3. Contexte marché et territoire", [
                    ("p", "Le Maroc bénéficie d'une dynamique touristique forte. Les chiffres publics du Ministère du Tourisme indiquent 19,8 millions de visiteurs en 2025, une progression de 14% par rapport à 2024, ainsi que 138 milliards de dirhams de recettes touristiques en 2025."),
                    ("p", "Tiznit offre une proposition distincte des destinations très saturées : médina, remparts, artisanat d'argent, proximité d'Aglou, accès aux paysages du Souss Massa et positionnement de séjour paisible."),
                    ("b", "Le rooftop permet de créer une offre visible au-delà de l'hébergement : café, restauration légère, soirées, privatisations et rendez-vous locaux."),
                ]),
                ("4. Atouts spécifiques de Mimosas", "advantages_table"),
                ("5. Dossier visuel", "visuals"),
                ("6. Business plan synthétique", "business_table"),
                ("7. Plan de reprise en 90 jours", "roadmap_table"),
                ("8. Argumentaire de vente", [
                    ("callout", "Pitch", "Hôtel Mimosas est une opportunité de reprise d'un actif hôtelier à fort potentiel à Tiznit : un projet déjà identifiable, doté d'un rooftop, d'espaces restaurant, d'une identité visuelle, d'un site multilingue, d'une visite virtuelle 360 et d'un storytelling territorial. Le repreneur ne part pas d'une page blanche : il reprend une base avancée, finalise les derniers travaux, structure l'exploitation et capte la valeur d'un hôtel indépendant dans une ville patrimoniale entre médina, Anti-Atlas et littoral d'Aglou."),
                    ("b", "Actif tangible : photos, espaces et identité déjà visibles."),
                    ("b", "Potentiel d'exploitation multiple : chambres, rooftop, restaurant, événements, expériences locales."),
                    ("b", "Outils commerciaux prêts : site multilingue, cartes, points d'intérêt, galerie et visite 360."),
                ]),
                ("9. Points de vigilance pour l'acheteur", [
                    ("b", "Confirmer le nombre exact de chambres exploitables, surfaces, capacité rooftop et capacité restaurant."),
                    ("b", "Chiffrer le CAPEX restant par lot : sécurité, plomberie, électricité, mobilier, cuisine, enseigne, accessibilité, classement."),
                    ("b", "Vérifier les autorisations, licences, conformité ERP/hôtel, assurance, fiscalité et statut foncier."),
                ]),
                ("10. Conclusion", [
                    ("p", "La valeur de Mimosas se joue dans la conversion d'un projet avancé en actif exploité. Le repreneur idéal n'achète pas seulement des murs ou des finitions : il achète une adresse, une marque, des espaces à potentiel, un rooftop différenciant, une base digitale et un récit touristique local prêt à être renforcé."),
                ]),
            ],
            "asset_rows": [
                ("Nom commercial", "Hôtel Mimosas"),
                ("Adresse", address["fr"]),
                ("Coordonnées GPS", f"{project['lat']}, {project['lng']}"),
                ("Concept", "Hôtel urbain avec rooftop, espaces restauration/bar et potentiel séjour court ou étape touristique."),
                ("État du projet", "Actif à potentiel avancé : image de marque, espaces visibles, contenus photos, visite virtuelle et site web déjà préparés. Les travaux, conformité et classement restent à confirmer par audit technique."),
                ("URL visite 360", project.get("tour_url", "Mimosas/Tour/index.htm")),
            ],
        },
        "en": {
            "title": "Hotel Mimosas",
            "subtitle": "Investor report, concise business plan and sale pitch",
            "meta": f"{address['en']} | Hotel asset to complete and operate",
            "caption_cover": "Facade and visual identity of Hotel Mimosas",
            "labels": ["Address", "Positioning", "Digital", "POI"],
            "kpis": [address["en"], "Urban hotel + rooftop + restaurant", "Multilingual website, Leaflet maps, gallery and 360 tour", f"{project.get('poi_count', 0)} mapped points of interest"],
            "sections": [
                ("1. Executive summary", [
                    ("callout", "Investment thesis", "Mimosas is a hotel asset in Tiznit whose value is already visible: brand identity, furnished spaces, rooftop, restaurant, 360 tour, photo gallery, multilingual website and an urban location close to the medina flows. The buyer can acquire an advanced project, complete the final finishes and launch an operation focused on stays, restaurant, rooftop and local experience."),
                    ("p", "This is not positioned as an abstract real-estate promise, but as a hotel operating platform. A buyer should immediately understand the path: finish, classify, commercialize, open, measure and optimize."),
                ]),
                ("2. Asset profile", "asset_table"),
                ("3. Market and destination context", [
                    ("p", "Morocco is benefiting from strong tourism momentum. Public figures from the Ministry of Tourism report 19.8 million visitors in 2025, up 14% versus 2024, and tourism receipts of MAD 138 billion in 2025."),
                    ("p", "Tiznit offers a distinctive alternative to saturated destinations: medina, ramparts, silver craft, proximity to Aglou, access to Souss Massa landscapes and a calm short-stay positioning."),
                    ("b", "The rooftop creates revenue beyond accommodation: cafe, light dining, evenings, private events and local meetups."),
                ]),
                ("4. Mimosas strengths", "advantages_table"),
                ("5. Visual dossier", "visuals"),
                ("6. Concise business plan", "business_table"),
                ("7. 90-day takeover plan", "roadmap_table"),
                ("8. Sale pitch", [
                    ("callout", "Pitch", "Hotel Mimosas is a takeover opportunity for a high-potential hotel asset in Tiznit: an identifiable project with a rooftop, restaurant spaces, visual identity, multilingual website, 360 virtual tour and destination storytelling. The buyer does not start from zero: they acquire an advanced base, complete the final works, structure the operation and capture the value of an independent hotel in a heritage city between the medina, the Anti-Atlas and Aglou's coastline."),
                    ("b", "Tangible asset: photos, spaces and identity are already visible."),
                    ("b", "Multiple operating potential: rooms, rooftop, restaurant, events and local experiences."),
                    ("b", "Commercial tools are ready: multilingual site, maps, points of interest, gallery and 360 tour."),
                ]),
                ("9. Buyer due diligence points", [
                    ("b", "Confirm the exact number of operable rooms, surfaces, rooftop capacity and restaurant capacity."),
                    ("b", "Cost the remaining CAPEX by trade: safety, plumbing, electricity, furniture, kitchen, signage, accessibility and classification."),
                    ("b", "Verify permits, licenses, hotel compliance, insurance, taxation and land/legal status."),
                ]),
                ("10. Conclusion", [
                    ("p", "Mimosas' value lies in converting an advanced project into an operating asset. The ideal buyer is not simply purchasing walls or finishes: they are buying an address, a brand, high-potential spaces, a differentiated rooftop, a digital base and a local tourism story ready to be strengthened."),
                ]),
            ],
            "asset_rows": [
                ("Commercial name", "Hotel Mimosas"),
                ("Address", address["en"]),
                ("GPS coordinates", f"{project['lat']}, {project['lng']}"),
                ("Concept", "Urban hotel with rooftop, food and beverage spaces and short-stay or tourist-stopover potential."),
                ("Project status", "Advanced potential asset: brand identity, visible spaces, photo content, virtual tour and website already prepared. Works, compliance and classification must be confirmed by technical audit."),
                ("360 tour URL", project.get("tour_url", "Mimosas/Tour/index.htm")),
            ],
        },
        "ar": {
            "title": "فندق ميموزاس",
            "subtitle": "تقرير للمستثمرين، خطة عمل مختصرة وحجة بيع",
            "meta": f"{address['ar']} | أصل فندقي جاهز لاستكمال التشطيبات والاستغلال",
            "caption_cover": "واجهة وهوية فندق ميموزاس",
            "labels": ["العنوان", "التموقع", "الرقمي", "نقاط الاهتمام"],
            "kpis": [address["ar"], "فندق حضري + سطح + مطعم", "موقع متعدد اللغات، خرائط Leaflet، معرض صور وجولة 360", f"{project.get('poi_count', 0)} نقطة اهتمام مدمجة في الخريطة"],
            "sections": [
                ("1. ملخص تنفيذي", [
                    ("callout", "فرضية الاستثمار", "ميموزاس أصل فندقي في تزنيت تظهر قيمته منذ الآن: هوية تجارية، فضاءات مهيأة، سطح، مطعم، جولة 360، معرض صور، موقع متعدد اللغات وموقع حضري قريب من حركة المدينة القديمة. الفرصة أمام المشتري هي اقتناء مشروع متقدم، إتمام اللمسات الأخيرة، ثم إطلاق الاستغلال الفندقي حول الإقامة والمطعم والسطح والتجربة المحلية."),
                    ("p", "لا يتم تقديم المشروع كوعود عقارية مجردة، بل كمنصة استغلال فندقي. يستطيع المستثمر تصور المسار بسرعة: الإنهاء، التصنيف، التسويق، الافتتاح، القياس والتحسين."),
                ]),
                ("2. بطاقة الأصل", "asset_table"),
                ("3. السوق والوجهة", [
                    ("p", "يعرف المغرب دينامية سياحية قوية. تشير أرقام وزارة السياحة إلى 19.8 مليون زائر سنة 2025، بزيادة 14% مقارنة بسنة 2024، وإلى 138 مليار درهم من مداخيل السياحة سنة 2025."),
                    ("p", "تزنيت تقدم عرضا مختلفا عن الوجهات المزدحمة: المدينة القديمة، الأسوار، صياغة الفضة، القرب من أكلو، وموقع هادئ ضمن سوس ماسة."),
                    ("b", "السطح يمكن أن يخلق مداخيل تتجاوز المبيت: مقهى، أكل خفيف، أمسيات، مناسبات خاصة ولقاءات محلية."),
                ]),
                ("4. نقاط قوة ميموزاس", "advantages_table"),
                ("5. الملف البصري", "visuals"),
                ("6. خطة عمل مختصرة", "business_table"),
                ("7. خطة استلام خلال 90 يوما", "roadmap_table"),
                ("8. حجة البيع", [
                    ("callout", "العرض المختصر", "فندق ميموزاس فرصة لاقتناء أصل فندقي واعد في تزنيت: مشروع معروف الهوية، يضم سطحا، فضاءات مطعم، هوية بصرية، موقعا متعدد اللغات، جولة افتراضية 360 وسردا سياحيا محليا. المشتري لا يبدأ من الصفر، بل يستلم قاعدة متقدمة، يكمل الأشغال الأخيرة، ينظم الاستغلال ويستفيد من قيمة فندق مستقل في مدينة تراثية بين المدينة القديمة والأطلس الصغير وساحل أكلو."),
                    ("b", "أصل ملموس: الصور والفضاءات والهوية موجودة وواضحة."),
                    ("b", "إمكانيات استغلال متعددة: غرف، سطح، مطعم، مناسبات وتجارب محلية."),
                    ("b", "أدوات تجارية جاهزة: موقع متعدد اللغات، خرائط، نقاط اهتمام، معرض صور وجولة 360."),
                ]),
                ("9. نقاط تحقق للمشتري", [
                    ("b", "تأكيد عدد الغرف القابلة للاستغلال والمساحات وطاقة السطح والمطعم."),
                    ("b", "تحديد كلفة الأشغال المتبقية: السلامة، السباكة، الكهرباء، الأثاث، المطبخ، اللافتات، الولوجيات والتصنيف."),
                    ("b", "التحقق من الرخص، المطابقة الفندقية، التأمين، الوضع الضريبي والقانوني."),
                ]),
                ("10. خلاصة", [
                    ("p", "قيمة ميموزاس تكمن في تحويل مشروع متقدم إلى أصل فندقي مستغل. المشتري المثالي لا يقتني الجدران والتشطيبات فقط، بل يقتني عنوانا، علامة، فضاءات واعدة، سطحا مميزا، قاعدة رقمية وحكاية سياحية محلية قابلة للتطوير."),
                ]),
            ],
            "asset_rows": [
                ("الاسم التجاري", "فندق ميموزاس"),
                ("العنوان", address["ar"]),
                ("الإحداثيات", f"{project['lat']}, {project['lng']}"),
                ("المفهوم", "فندق حضري مع سطح وفضاءات مطعم وإمكانية إقامة قصيرة أو سياحية."),
                ("حالة المشروع", "أصل متقدم الإمكانيات: هوية تجارية، فضاءات ظاهرة، صور، جولة افتراضية وموقع جاهز. تبقى الأشغال والمطابقة والتصنيف رهينة بتدقيق تقني."),
                ("رابط جولة 360", project.get("tour_url", "Mimosas/Tour/index.htm")),
            ],
        },
        "es": {
            "title": "Hotel Mimosas",
            "subtitle": "Informe para inversores, plan de negocio sintético y argumento de venta",
            "meta": f"{address['es']} | Activo hotelero para terminar y explotar",
            "caption_cover": "Fachada e identidad visual del Hotel Mimosas",
            "labels": ["Dirección", "Posicionamiento", "Digital", "POI"],
            "kpis": [address["es"], "Hotel urbano + rooftop + restaurante", "Sitio multilingüe, mapas Leaflet, galería y visita 360", f"{project.get('poi_count', 0)} puntos de interés cartografiados"],
            "sections": [
                ("1. Resumen ejecutivo", [
                    ("callout", "Tesis de inversión", "Mimosas es un activo hotelero en Tiznit cuyo valor ya es visible: identidad de marca, espacios acondicionados, rooftop, restaurante, visita 360, galería de fotos, sitio multilingüe y ubicación urbana cercana a la medina. La oportunidad para el comprador es adquirir un proyecto avanzado, ejecutar los últimos acabados y lanzar una explotación orientada a estancia, restaurante, rooftop y experiencia local."),
                    ("p", "El proyecto no se presenta como una promesa inmobiliaria abstracta, sino como una plataforma de explotación hotelera. El comprador puede proyectarse rápidamente: terminar, clasificar, comercializar, abrir, medir y optimizar."),
                ]),
                ("2. Ficha del activo", "asset_table"),
                ("3. Mercado y territorio", [
                    ("p", "Marruecos disfruta de una fuerte dinámica turística. Las cifras públicas del Ministerio de Turismo indican 19,8 millones de visitantes en 2025, un crecimiento del 14% frente a 2024, y 138 mil millones de dirhams de ingresos turísticos en 2025."),
                    ("p", "Tiznit ofrece una propuesta distinta de los destinos saturados: medina, murallas, artesanía de plata, proximidad a Aglou, acceso a paisajes de Souss Massa y posicionamiento de estancia tranquila."),
                    ("b", "El rooftop permite crear ingresos más allá del alojamiento: café, comida ligera, tardes, eventos privados y encuentros locales."),
                ]),
                ("4. Fortalezas de Mimosas", "advantages_table"),
                ("5. Dossier visual", "visuals"),
                ("6. Plan de negocio sintético", "business_table"),
                ("7. Plan de toma de control en 90 días", "roadmap_table"),
                ("8. Argumento de venta", [
                    ("callout", "Pitch", "Hotel Mimosas es una oportunidad de adquisición de un activo hotelero con alto potencial en Tiznit: un proyecto ya identificable, con rooftop, espacios de restaurante, identidad visual, sitio multilingüe, visita virtual 360 y narrativa territorial. El comprador no parte de cero: adquiere una base avanzada, termina las últimas obras, estructura la explotación y captura el valor de un hotel independiente en una ciudad patrimonial entre la medina, el Anti-Atlas y la costa de Aglou."),
                    ("b", "Activo tangible: fotos, espacios e identidad ya visibles."),
                    ("b", "Potencial de explotación múltiple: habitaciones, rooftop, restaurante, eventos y experiencias locales."),
                    ("b", "Herramientas comerciales listas: sitio multilingüe, mapas, puntos de interés, galería y visita 360."),
                ]),
                ("9. Puntos de verificación para el comprador", [
                    ("b", "Confirmar el número exacto de habitaciones explotables, superficies, capacidad del rooftop y del restaurante."),
                    ("b", "Valorar el CAPEX restante: seguridad, fontanería, electricidad, mobiliario, cocina, señalética, accesibilidad y clasificación."),
                    ("b", "Verificar autorizaciones, licencias, cumplimiento hotelero, seguros, fiscalidad y situación jurídica."),
                ]),
                ("10. Conclusión", [
                    ("p", "El valor de Mimosas está en convertir un proyecto avanzado en un activo explotado. El comprador ideal no compra solo paredes o acabados: compra una dirección, una marca, espacios con potencial, un rooftop diferenciado, una base digital y un relato turístico local listo para reforzarse."),
                ]),
            ],
            "asset_rows": [
                ("Nombre comercial", "Hotel Mimosas"),
                ("Dirección", address["es"]),
                ("Coordenadas GPS", f"{project['lat']}, {project['lng']}"),
                ("Concepto", "Hotel urbano con rooftop, espacios de restaurante/bar y potencial para estancias cortas o escala turística."),
                ("Estado del proyecto", "Activo avanzado con potencial: identidad de marca, espacios visibles, fotos, visita virtual y sitio web ya preparados. Obras, conformidad y clasificación deben confirmarse por auditoría técnica."),
                ("URL visita 360", project.get("tour_url", "Mimosas/Tour/index.htm")),
            ],
        },
    }


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="E5E7EB", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_run(p, text, bold=False, italic=False, color=None, size=10):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Arial" if any("\u0600" <= ch <= "\u06ff" for ch in text) else "Aptos"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def add_heading(doc, text, lang, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, text, bold=True, color=PALETTE["brown"] if level == 1 else PALETTE["ink"], size=18 if level == 1 else 13)


def add_body(doc, text, lang):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    add_run(p, text, color=PALETTE["ink"], size=10.2)


def add_bullet(doc, text, lang):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, color=PALETTE["ink"], size=9.8)


def shade_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, PALETTE["brown"])
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True


def add_callout(doc, title, body, lang):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALETTE["cream"])
    set_cell_border(cell, "D7B56D", "10")
    set_cell_margins(cell, 170, 220, 170, 220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, title, bold=True, color=PALETTE["brown"], size=11.2)
    p2 = cell.add_paragraph()
    p2.alignment = p.alignment
    add_run(p2, body, color=PALETTE["ink"], size=9.6)


def add_photo(doc, path, caption, lang, width=6.2):
    if not path:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, color=PALETTE["muted"], size=8.4)


def add_photo_grid(doc, lang):
    captions = {
        "fr": ["Devanture", "Entrée", "Réception", "Chambre", "Couloir", "Salle d'eau", "Rooftop", "Restaurant"],
        "en": ["Facade", "Entrance", "Reception", "Room", "Corridor", "Bathroom", "Rooftop", "Restaurant"],
        "ar": ["الواجهة", "المدخل", "الاستقبال", "الغرفة", "الممر", "الحمام", "السطح", "المطعم"],
        "es": ["Fachada", "Entrada", "Recepción", "Habitación", "Pasillo", "Baño", "Rooftop", "Restaurante"],
    }[lang]
    keys = ["facade", "entree", "reception", "chambre", "couloir", "sdb", "rooftop", "restaurant"]
    table = doc.add_table(rows=0, cols=2)
    for idx in range(0, len(keys), 2):
        row = table.add_row()
        for col in range(2):
            key_idx = idx + col
            cell = row.cells[col]
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, 80, 80, 120, 80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(IMAGES[keys[key_idx]]), width=Inches(3.05))
            cap = cell.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(cap, captions[key_idx], italic=True, color=PALETTE["muted"], size=8)


def add_kpi_table(doc, data, lang):
    table = doc.add_table(rows=1, cols=4)
    for i, label in enumerate(data["labels"]):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, PALETTE["cream"])
        set_cell_border(cell, "E6D3AA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label, bold=True, color=PALETTE["brown"], size=8.3)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, data["kpis"][i], color=PALETTE["ink"], size=9)


def add_two_col_table(doc, rows, headers, lang):
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = headers[0]
    table.rows[0].cells[1].text = headers[1]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    shade_table(table)


def add_three_col_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    shade_table(table)


ADV = {
    "fr": (["Atout", "Impact commercial", "Action recommandée"], [
        ("Rooftop", "Différenciation forte et revenus hors hébergement.", "Créer une offre sunset, petit-déjeuner, café, tapas et événements privés."),
        ("Restaurant / comptoir", "Revenus additionnels et meilleure expérience client.", "Définir une carte courte, locale, rentable et facile à exécuter."),
        ("Visite 360", "Réduit l'incertitude investisseur et client.", "L'utiliser dans la vente, Google Business, le site et les annonces."),
    ]),
    "en": (["Strength", "Commercial impact", "Recommended action"], [
        ("Rooftop", "Strong differentiation and revenue beyond rooms.", "Create sunset, breakfast, cafe, tapas and private-event offers."),
        ("Restaurant / counter", "Additional revenue and stronger guest experience.", "Define a short, local, profitable and easy-to-run menu."),
        ("360 tour", "Reduces investor and guest uncertainty.", "Use it in the sale process, Google Business, the website and listings."),
    ]),
    "ar": (["نقطة قوة", "الأثر التجاري", "الإجراء المقترح"], [
        ("السطح", "تمييز قوي ومداخيل خارج الإقامة.", "إطلاق عروض الغروب، الفطور، المقهى، الأكل الخفيف والمناسبات الخاصة."),
        ("المطعم / الكاونتر", "مداخيل إضافية وتجربة أفضل للضيف.", "اعتماد قائمة قصيرة ومحلية ومربحة وسهلة التنفيذ."),
        ("جولة 360", "تقلل تردد المستثمر والزبون.", "استعمالها في البيع وGoogle Business والموقع والإعلانات."),
    ]),
    "es": (["Fortaleza", "Impacto comercial", "Acción recomendada"], [
        ("Rooftop", "Fuerte diferenciación e ingresos fuera del alojamiento.", "Crear ofertas de sunset, desayuno, café, tapas y eventos privados."),
        ("Restaurante / barra", "Ingresos adicionales y mejor experiencia cliente.", "Definir una carta corta, local, rentable y fácil de ejecutar."),
        ("Visita 360", "Reduce la incertidumbre de inversores y clientes.", "Usarla en la venta, Google Business, el sitio y los anuncios."),
    ]),
}

BUS = {
    "fr": (["Pilier", "Hypothèse opérationnelle", "KPI"], [("Hébergement", "Vente directe + OTA + clientèle locale, MRE, voyageurs d'étape.", "Occupation, ADR, RevPAR"), ("Rooftop", "Café, brunch, sunset, tapas, événements privés.", "Ticket moyen, couverts, marge"), ("Digital", "Site, Google Business, WhatsApp, visite 360, SEO local.", "Trafic, demandes, conversion")]),
    "en": (["Pillar", "Operating assumption", "KPI"], [("Accommodation", "Direct sales + OTAs + local guests, diaspora, stopover travelers.", "Occupancy, ADR, RevPAR"), ("Rooftop", "Cafe, brunch, sunset, tapas, private events.", "Average ticket, covers, margin"), ("Digital", "Website, Google Business, WhatsApp, 360 tour, local SEO.", "Traffic, enquiries, conversion")]),
    "ar": (["المحور", "فرضية التشغيل", "المؤشر"], [("الإقامة", "بيع مباشر + منصات + زبائن محليون ومغاربة العالم ومسافرون.", "الإشغال، السعر، RevPAR"), ("السطح", "مقهى، فطور متأخر، غروب، أكل خفيف، مناسبات.", "متوسط الفاتورة، الزبائن، الهامش"), ("الرقمي", "موقع، Google Business، واتساب، جولة 360، SEO محلي.", "الزيارات، الطلبات، التحويل")]),
    "es": (["Pilar", "Hipótesis operativa", "KPI"], [("Alojamiento", "Venta directa + OTAs + clientes locales, diáspora, viajeros de paso.", "Ocupación, ADR, RevPAR"), ("Rooftop", "Café, brunch, sunset, tapas, eventos privados.", "Ticket medio, cubiertos, margen"), ("Digital", "Sitio, Google Business, WhatsApp, visita 360, SEO local.", "Tráfico, solicitudes, conversión")]),
}

ROAD = {
    "fr": (["Phase", "Objectif", "Actions"], [("Jours 1-15", "Audit et sécurisation", "Audit technique, juridique, conformité, sécurité, licences."), ("Jours 16-45", "Finitions", "Chambres, éclairage, plomberie, rooftop, cuisine, réception."), ("Jours 46-90", "Commercialisation et soft opening", "Photos finales, OTA, Google Business, site, avis clients.")]),
    "en": (["Phase", "Goal", "Actions"], [("Days 1-15", "Audit and secure", "Technical, legal, compliance, safety and license audit."), ("Days 16-45", "Final finishes", "Rooms, lighting, plumbing, rooftop, kitchen, reception."), ("Days 46-90", "Marketing and soft opening", "Final photos, OTAs, Google Business, website, reviews.")]),
    "ar": (["المرحلة", "الهدف", "الإجراءات"], [("الأيام 1-15", "التدقيق والتأمين", "تدقيق تقني وقانوني ومطابقة وسلامة ورخص."), ("الأيام 16-45", "التشطيبات", "الغرف، الإنارة، السباكة، السطح، المطبخ، الاستقبال."), ("الأيام 46-90", "التسويق والافتتاح التجريبي", "صور نهائية، منصات، Google Business، الموقع، الآراء.")]),
    "es": (["Fase", "Objetivo", "Acciones"], [("Días 1-15", "Auditoría y seguridad", "Auditoría técnica, legal, conformidad, seguridad y licencias."), ("Días 16-45", "Acabados finales", "Habitaciones, iluminación, fontanería, rooftop, cocina, recepción."), ("Días 46-90", "Comercialización y soft opening", "Fotos finales, OTAs, Google Business, sitio, reseñas.")]),
}


def build(lang, data):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.45)
        s.bottom_margin = Cm(1.45)
        s.left_margin = Cm(1.55)
        s.right_margin = Cm(1.55)
    doc.styles["Normal"].font.name = "Arial" if lang == "ar" else "Aptos"
    doc.styles["Normal"].font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(IMAGES["logo"]), width=Inches(2.1))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, data["title"], bold=True, color=PALETTE["brown"], size=29)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, data["subtitle"], color=PALETTE["ink"], size=13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, data["meta"], color=PALETTE["muted"], size=10)
    add_photo(doc, IMAGES["facade"], data["caption_cover"], lang, 6.65)
    add_kpi_table(doc, data, lang)
    doc.add_page_break()

    for title, content in data["sections"]:
        add_heading(doc, title, lang)
        if content == "asset_table":
            headers = {"fr": ("Critère", "Lecture investisseur"), "en": ("Criterion", "Investor reading"), "ar": ("المعيار", "قراءة المستثمر"), "es": ("Criterio", "Lectura inversor")}[lang]
            add_two_col_table(doc, data["asset_rows"], headers, lang)
        elif content == "advantages_table":
            headers, rows = ADV[lang]
            add_three_col_table(doc, headers, rows)
        elif content == "business_table":
            add_callout(doc, {"fr": "Important", "en": "Important", "ar": "مهم", "es": "Importante"}[lang], {
                "fr": "Les chiffres précis doivent être finalisés après audit : prix, nombre de chambres, CAPEX restant, autorisations et politique tarifaire.",
                "en": "Precise figures must be finalized after audit: price, room count, remaining CAPEX, permits and pricing policy.",
                "ar": "يجب تثبيت الأرقام الدقيقة بعد التدقيق: السعر، عدد الغرف، الكلفة المتبقية، الرخص وسياسة الأسعار.",
                "es": "Las cifras precisas deben finalizarse tras auditoría: precio, número de habitaciones, CAPEX restante, permisos y política tarifaria.",
            }[lang], lang)
            headers, rows = BUS[lang]
            add_three_col_table(doc, headers, rows)
        elif content == "roadmap_table":
            headers, rows = ROAD[lang]
            add_three_col_table(doc, headers, rows)
        elif content == "visuals":
            add_body(doc, {
                "fr": "Les visuels montrent un actif déjà lisible : identité, façade, accueil, chambres, rooftop et restaurant.",
                "en": "The visuals show an already legible asset: identity, facade, reception, rooms, rooftop and restaurant.",
                "ar": "تظهر الصور أصلا واضحا منذ الآن: الهوية، الواجهة، الاستقبال، الغرف، السطح والمطعم.",
                "es": "Los visuales muestran un activo ya legible: identidad, fachada, recepción, habitaciones, rooftop y restaurante.",
            }[lang], lang)
            add_photo_grid(doc, lang)
        else:
            for item in content:
                kind = item[0]
                if kind == "p":
                    add_body(doc, item[1], lang)
                elif kind == "b":
                    add_bullet(doc, item[1], lang)
                elif kind == "callout":
                    add_callout(doc, item[1], item[2], lang)

    add_heading(doc, {"fr": "Sources", "en": "Sources", "ar": "المصادر", "es": "Fuentes"}[lang], lang)
    for source in [
        "Ministère du Tourisme du Maroc - chiffres clés 2025",
        "Visit Agadir / Souss Massa - Tiznit et Aglou",
        "Données internes Mimosas - galerie, site, visite 360, data/projects.json",
    ]:
        add_bullet(doc, source, lang)
    out = OUT_DIR / f"rapport-investisseur-hotel-mimosas-{lang}.docx"
    doc.core_properties.title = data["title"]
    doc.core_properties.subject = data["subtitle"]
    doc.core_properties.author = "Hotel Mimosas"
    doc.save(out)
    return out


for lang, data in tr_text().items():
    print(build(lang, data))
